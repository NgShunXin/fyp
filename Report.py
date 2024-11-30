from flask import Flask, render_template, request, send_file, jsonify
from jamaibase import JamAI, protocol as p
import json
from docx import Document
from io import BytesIO
import random
import string
import os

app = Flask(__name__)

# Initialize JamAI
jamai = JamAI(api_key="jamai_sk_09081aedfccde72a8cfa4bc4db0ff23fa5bd47406c885b77", project_id="proj_2e6b08f124289d82c1e430d3")

# Generate random filename
def generate_random_filename(extension=".docx"):
    random_str = ''.join(random.choices(string.ascii_letters + string.digits, k=10))
    return f"final_report_{random_str}{extension}"

@app.route('/', methods=['GET'])
def index():
    error = None
    try:
        # Load test history data from the file
        with open('static/test_history.json', 'r') as file:
            test_history = json.load(file)
            # Convert the entire test_history to a formatted string
            test_history_text = json.dumps(test_history, indent=2)

        # Add rows to the existing table with the text data
        completion = jamai.add_table_rows(
            "action",
            p.RowAddRequest(
                table_id="report",
                data=[{"input": test_history_text}],  # Pass the text version instead of JSON object
                stream=False
            )
        )

        if completion.rows:
            output_row = completion.rows[0].columns
            # Extract text content from each field
            summary = output_row.get("Summary", {}).text if output_row.get("Summary") else 'N/A'
            performances = output_row.get("Performance", {}).text if output_row.get("Performance") else 'N/A'
            trends = output_row.get("Trends", {}).text if output_row.get("Trends") else 'N/A'
            recommendations = output_row.get("Recommendations", {}).text if output_row.get("Recommendations") else 'N/A'

            # Generate report as DOCX
            doc = Document()
            doc.add_heading("Executive Report", level=1)

            doc.add_heading("Summary of Test History", level=2)
            doc.add_paragraph(summary)

            doc.add_heading("Emotion Recognition Performance", level=2)
            doc.add_paragraph(performances)

            doc.add_heading("Key Insights and Trends", level=2)
            doc.add_paragraph(trends)

            doc.add_heading("Recommendations for Improvement", level=2)
            doc.add_paragraph(recommendations)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return send_file(
                buffer,
                as_attachment=True,
                download_name=generate_random_filename(),
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            error = "⚠️ Failed to get a response. Please try again."

    except Exception as e:
        error = f"❌ An error occurred: {e}"
        print(f"Error details: {str(e)}")  # For debugging

    return render_template("report.html", error=error)

if __name__ == '__main__':
    app.run(debug=True)


